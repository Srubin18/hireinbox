#!/usr/bin/env python3
"""Convert Markdown PRD to Word document"""

import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_hyperlink(paragraph, url, text):
    """Add a hyperlink to a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    c = OxmlElement('w:color')
    c.set(qn('w:val'), '0563C1')
    rPr.append(c)
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink

def set_cell_shading(cell, color):
    """Set cell background color."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color)
    tcPr.append(shd)

def parse_markdown_to_docx(md_content, output_path):
    """Parse markdown and create Word document."""
    doc = Document()

    # Set up styles
    styles = doc.styles

    # Title style
    title_style = styles['Title']
    title_style.font.size = Pt(28)
    title_style.font.bold = True
    title_style.font.color.rgb = RGBColor(0, 51, 102)

    # Heading styles
    for i in range(1, 5):
        h_style = styles[f'Heading {i}']
        h_style.font.color.rgb = RGBColor(0, 51, 102)
        h_style.font.bold = True
        if i == 1:
            h_style.font.size = Pt(20)
        elif i == 2:
            h_style.font.size = Pt(16)
        elif i == 3:
            h_style.font.size = Pt(14)
        else:
            h_style.font.size = Pt(12)

    lines = md_content.split('\n')
    i = 0
    in_code_block = False
    code_buffer = []
    in_table = False
    table_rows = []

    while i < len(lines):
        line = lines[i]

        # Skip empty lines
        if not line.strip():
            if in_table and table_rows:
                # End table
                create_table(doc, table_rows)
                table_rows = []
                in_table = False
            i += 1
            continue

        # Horizontal rule
        if line.strip() == '---':
            if in_table and table_rows:
                create_table(doc, table_rows)
                table_rows = []
                in_table = False
            doc.add_paragraph('_' * 50)
            i += 1
            continue

        # Code block
        if line.strip().startswith('```'):
            if in_code_block:
                # End code block
                code_text = '\n'.join(code_buffer)
                p = doc.add_paragraph()
                p.style = 'No Spacing'
                run = p.add_run(code_text)
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
                code_buffer = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_buffer.append(line)
            i += 1
            continue

        # Table row
        if '|' in line and line.strip().startswith('|'):
            # Check if this is a separator row
            if re.match(r'^[\|\-\s:]+$', line.strip()):
                i += 1
                continue

            in_table = True
            cells = [c.strip() for c in line.split('|')[1:-1]]
            table_rows.append(cells)
            i += 1
            continue

        # If we were in a table, create it
        if in_table and table_rows:
            create_table(doc, table_rows)
            table_rows = []
            in_table = False

        # Headers
        if line.startswith('# '):
            text = line[2:].strip()
            if 'HIREINBOX' in text and 'Technical Product' in text:
                doc.add_heading(text, 0)
            else:
                doc.add_heading(text, 1)
            i += 1
            continue

        if line.startswith('## '):
            doc.add_heading(line[3:].strip(), 2)
            i += 1
            continue

        if line.startswith('### '):
            doc.add_heading(line[4:].strip(), 3)
            i += 1
            continue

        if line.startswith('#### '):
            doc.add_heading(line[5:].strip(), 4)
            i += 1
            continue

        # Bullet points
        if line.strip().startswith('- ') or line.strip().startswith('* '):
            text = line.strip()[2:]
            text = process_inline_formatting(text)
            p = doc.add_paragraph(text, style='List Bullet')
            i += 1
            continue

        # Numbered list
        if re.match(r'^\d+\.\s', line.strip()):
            text = re.sub(r'^\d+\.\s', '', line.strip())
            text = process_inline_formatting(text)
            p = doc.add_paragraph(text, style='List Number')
            i += 1
            continue

        # Regular paragraph
        text = process_inline_formatting(line)
        if text.strip():
            doc.add_paragraph(text)

        i += 1

    # Handle any remaining table
    if table_rows:
        create_table(doc, table_rows)

    doc.save(output_path)
    print(f"Document saved to: {output_path}")

def process_inline_formatting(text):
    """Process inline markdown formatting."""
    # Remove markdown links but keep text
    text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', text)
    # Remove bold markers
    text = re.sub(r'\*\*([^\*]+)\*\*', r'\1', text)
    # Remove italic markers
    text = re.sub(r'\*([^\*]+)\*', r'\1', text)
    # Remove inline code markers
    text = re.sub(r'`([^`]+)`', r'\1', text)
    return text

def create_table(doc, rows):
    """Create a table from parsed rows."""
    if not rows:
        return

    num_cols = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.style = 'Table Grid'

    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            if j < num_cols:
                cell = table.cell(i, j)
                cell.text = process_inline_formatting(cell_text)
                # Header row styling
                if i == 0:
                    set_cell_shading(cell, 'E6E6E6')
                    cell.paragraphs[0].runs[0].bold = True if cell.paragraphs[0].runs else False

    # Add spacing after table
    doc.add_paragraph()

def main():
    import sys

    files = [
        ('/Users/simon/Desktop/hireinbox/HIREINBOX_Technical_PRD.md',
         '/Users/simon/Desktop/hireinbox/HIREINBOX_Technical_PRD.docx'),
        ('/Users/simon/Desktop/hireinbox/HIREINBOX_Partner_PRD.md',
         '/Users/simon/Desktop/hireinbox/HIREINBOX_Partner_PRD.docx'),
        ('/Users/simon/Desktop/hireinbox/HIREINBOX_Business_PRD.md',
         '/Users/simon/Desktop/hireinbox/HIREINBOX_Business_PRD.docx'),
    ]

    for input_path, output_path in files:
        try:
            with open(input_path, 'r', encoding='utf-8') as f:
                md_content = f.read()
            parse_markdown_to_docx(md_content, output_path)
        except FileNotFoundError:
            print(f"Skipping {input_path} - file not found")

if __name__ == '__main__':
    main()
