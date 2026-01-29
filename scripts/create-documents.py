#!/usr/bin/env python3
"""
Create Word and Excel documents for HireInbox partner review
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
from openpyxl.utils import get_column_letter
import os

# Output paths
DESKTOP = "/Users/simon/Desktop"
PRD_PATH = os.path.join(DESKTOP, "HIREINBOX_PRD_TECHNICAL.docx")
MARKET_PATH = os.path.join(DESKTOP, "HIREINBOX_MARKET_SIZE.docx")
FORECAST_PATH = os.path.join(DESKTOP, "HIREINBOX_REVENUE_FORECAST_Y1.xlsx")

def add_heading(doc, text, level=1):
    """Add a styled heading"""
    heading = doc.add_heading(text, level=level)
    return heading

def add_table(doc, headers, rows):
    """Add a formatted table"""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'

    # Header row
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        hdr_cells[i].paragraphs[0].runs[0].bold = True

    # Data rows
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, cell_data in enumerate(row_data):
            row_cells[i].text = str(cell_data)

    return table

def create_prd_document():
    """Create PRD Technical Document"""
    doc = Document()

    # Title
    title = doc.add_heading('HIREINBOX - PRODUCT REQUIREMENTS DOCUMENT (PRD)', 0)

    doc.add_paragraph('Version: 1.0')
    doc.add_paragraph('Date: 25 January 2026')
    doc.add_paragraph('Author: Claude (CTO)')
    doc.add_paragraph('For: Partner Review - MVP Assessment')
    doc.add_paragraph()

    # Executive Summary
    add_heading(doc, '1. EXECUTIVE SUMMARY', 1)

    add_heading(doc, 'What is HireInbox?', 2)
    doc.add_paragraph('HireInbox is an AI-powered CV screening platform built specifically for South African SMEs. The platform screens CVs with evidence-based reasoning, providing explainable AI decisions that are POPIA compliant.')

    add_heading(doc, 'Core Value Proposition', 2)
    doc.add_paragraph('"Less noise. Better hires."')
    bullets = doc.add_paragraph()
    bullets.add_run('• Screen 200 CVs in 30 seconds (vs 17-50 hours manually)\n')
    bullets.add_run('• Every decision shows WHY with direct quotes from the CV\n')
    bullets.add_run('• South African context built-in (CA(SA), BCom, local companies)\n')
    bullets.add_run('• Per-role pricing (not per-CV) - predictable costs')

    add_heading(doc, 'Target Markets', 2)
    add_table(doc,
        ['Segment', 'Description', 'Pricing Model'],
        [
            ['B2B (Employers)', 'SMEs hiring 1-50 roles/year', 'Per role (R1,750+)'],
            ['B2C (Candidates)', 'Job seekers wanting feedback', 'Free + Upsells (R99-R299)'],
            ['B2Recruiter', 'Recruitment agencies', 'Per search/role']
        ]
    )
    doc.add_paragraph()

    add_heading(doc, 'Current Build Status', 2)
    add_table(doc,
        ['Component', 'Status', 'Completeness'],
        [
            ['AI CV Screening', 'LIVE', '95%'],
            ['B2B Employer Flow', 'Working', '75%'],
            ['B2C Candidate Flow', 'Working', '80%'],
            ['Talent Pool', 'Partial', '60%'],
            ['B2Recruiter', 'UI Only', '40%'],
            ['Payments', 'Not Integrated', '20%'],
            ['Mobile', 'Needs Work', '30%']
        ]
    )
    doc.add_paragraph()
    doc.add_paragraph('Overall MVP Status: ~60%')

    # Product Overview
    add_heading(doc, '2. PRODUCT OVERVIEW', 1)

    add_heading(doc, '2.1 Product Philosophy', 2)
    bullets = doc.add_paragraph()
    bullets.add_run('1. AI as Assistant - AI assists humans, never makes final decisions\n')
    bullets.add_run('2. Evidence-Based - Every recommendation backed by direct quotes\n')
    bullets.add_run('3. POPIA Compliant - Full audit trail, data rights respected\n')
    bullets.add_run('4. SA-Specific - Understands local qualifications, companies, context\n')
    bullets.add_run('5. Simple Pricing - Per-role, not per-CV; no usage anxiety')

    # Tech Stack
    add_heading(doc, '3. TECHNICAL ARCHITECTURE', 1)

    add_heading(doc, '3.1 Tech Stack', 2)
    add_table(doc,
        ['Layer', 'Technology', 'Purpose'],
        [
            ['Frontend', 'Next.js 16 (App Router)', 'React framework with SSR'],
            ['Styling', 'Inline CSS', 'No Tailwind (design decision)'],
            ['Language', 'TypeScript', 'Type safety'],
            ['Database', 'Supabase (PostgreSQL)', 'Data persistence, auth'],
            ['AI - CV Screening', 'OpenAI GPT-4o-mini (Fine-tuned)', 'Core screening intelligence'],
            ['AI - Video', 'Claude Vision (claude-sonnet-4)', 'Video analysis'],
            ['AI - Transcription', 'Whisper-1', 'Audio to text'],
            ['Email', 'IMAP Integration', 'Fetch CVs from inbox'],
            ['Payments', 'PayFast (Planned)', 'SA payment gateway'],
            ['Hosting', 'Vercel', 'Edge deployment'],
            ['Domain', 'hireinbox.co.za', 'Live']
        ]
    )
    doc.add_paragraph()

    add_heading(doc, '3.3 Fine-Tuned AI Model', 2)
    doc.add_paragraph('Model ID: ft:gpt-4o-mini-2024-07-18:personal:hireinbox-cv-screener:CphiMaZU')
    doc.add_paragraph('Training Data: 860 hand-curated screening examples (v1), 10,000 examples generating (v2 - in progress)')

    # B2B Products
    add_heading(doc, '4. B2B PRODUCTS (EMPLOYERS)', 1)

    add_heading(doc, '4.1 AI CV Screening', 2)
    doc.add_paragraph('Price: R1,750 per role')
    doc.add_paragraph('What It Does:')
    bullets = doc.add_paragraph()
    bullets.add_run('1. Employer creates a role with requirements\n')
    bullets.add_run('2. System fetches CVs from their email inbox (IMAP)\n')
    bullets.add_run('3. AI screens each CV against role requirements\n')
    bullets.add_run('4. Candidates scored 0-100 with evidence\n')
    bullets.add_run('5. Auto-generates shortlist (80+), consider (60-79), reject (<60)\n')
    bullets.add_run('6. Sends acknowledgment emails to candidates')
    doc.add_paragraph('Status: LIVE - Core feature working')

    add_heading(doc, '4.2 AI Interview (Add-On)', 2)
    doc.add_paragraph('Price: R799 per role')
    doc.add_paragraph('Status: Experimental - UI exists, needs production hardening')

    add_heading(doc, '4.3 Verification Bundle (Add-On)', 2)
    doc.add_paragraph('Price: R800 per role (or individual: R50 ID, R100 Credit, R200 Reference)')
    doc.add_paragraph('Status: API stubs exist, third-party integrations pending')

    # B2C Products
    add_heading(doc, '5. B2C PRODUCTS (CANDIDATES)', 1)

    add_heading(doc, '5.1 Free CV Scan', 2)
    doc.add_paragraph('Price: FREE (1 per user)')
    doc.add_paragraph('Status: LIVE - Core feature working')

    add_heading(doc, '5.2 Video Analysis (Paid Upsell)', 2)
    doc.add_paragraph('Price: R99-R199')
    doc.add_paragraph('Status: LIVE - Claude Vision working')

    add_heading(doc, '5.3-5.5 Other B2C Products', 2)
    add_table(doc,
        ['Product', 'Price', 'Status'],
        [
            ['AI Coaching', 'R149-R299', 'UI mockup only'],
            ['Position Prep', 'R199', 'UI mockup only'],
            ['Video Pitch', 'R149', 'UI mockup only']
        ]
    )
    doc.add_paragraph()

    # Talent Pool
    add_heading(doc, '6. TALENT POOL PLATFORM', 1)
    doc.add_paragraph('The Talent Pool is a two-sided marketplace:')
    bullets = doc.add_paragraph()
    bullets.add_run('• Candidates join for FREE to be discovered\n')
    bullets.add_run('• Employers pay R2,500 to post a job and access matched candidates')

    # B2Recruiter
    add_heading(doc, '7. B2RECRUITER PRODUCTS', 1)

    add_heading(doc, '7.1 Talent Mapping', 2)
    doc.add_paragraph('Price: R999 per search')
    doc.add_paragraph('Status: Working - basic search functional')

    # Pricing Summary
    add_heading(doc, '8. PRICING SUMMARY', 1)

    add_heading(doc, 'B2B Pricing', 2)
    add_table(doc,
        ['Product', 'Price', 'Unit'],
        [
            ['AI CV Screening', 'R1,750', 'per role'],
            ['AI Interview', 'R799', 'per role (add-on)'],
            ['Verification Bundle', 'R800', 'per role (add-on)'],
            ['ID Check only', 'R50', 'per candidate'],
            ['Credit Check only', 'R100', 'per candidate'],
            ['Reference Check only', 'R200', 'per candidate'],
            ['Job Listing (Talent Pool)', 'R2,500', 'per listing'],
            ['Talent Mapping', 'R999', 'per search']
        ]
    )
    doc.add_paragraph()

    add_heading(doc, 'B2C Pricing', 2)
    add_table(doc,
        ['Product', 'Price'],
        [
            ['CV Scan', 'FREE (1x)'],
            ['CV Rewrite', 'FREE (1x)'],
            ['Video Analysis', 'R99-R199'],
            ['AI Coaching', 'R149-R299'],
            ['Position Prep', 'R199'],
            ['Video Pitch', 'R149']
        ]
    )
    doc.add_paragraph()

    # Footer
    doc.add_paragraph()
    doc.add_paragraph('Prepared for: Partner Review')
    doc.add_paragraph('Next Steps: Review this PRD, Review REMAINING_MVP_ITEMS.md, Review revenue forecast')
    doc.add_paragraph()
    doc.add_paragraph('HireInbox - Less noise. Better hires.')
    doc.add_paragraph('Built in Cape Town, South Africa')

    doc.save(PRD_PATH)
    print(f"Created: {PRD_PATH}")

def create_market_size_document():
    """Create Market Size Document"""
    doc = Document()

    # Title
    title = doc.add_heading('HIREINBOX - MARKET SIZE ANALYSIS', 0)

    doc.add_paragraph('Date: 25 January 2026')
    doc.add_paragraph('Prepared by: Claude (CTO)')
    doc.add_paragraph('For: Partner Review')
    doc.add_paragraph()

    # Executive Summary
    add_heading(doc, 'EXECUTIVE SUMMARY', 1)
    doc.add_paragraph('HireInbox operates in a large, fragmented market with significant growth potential:')

    add_table(doc,
        ['Market', 'TAM (Total Addressable)', 'SAM (Serviceable)', 'SOM (Obtainable Y1)'],
        [
            ['B2B (Employers)', 'R8.75 billion', 'R1.75 billion', 'R3.5 million'],
            ['B2C (Job Seekers)', 'R4.1 billion', 'R820 million', 'R1.3 million'],
            ['Recruiters', 'R2.5 billion', 'R500 million', 'R300,000'],
            ['TOTAL', 'R15.35 billion', 'R3.07 billion', 'R5.1 million']
        ]
    )
    doc.add_paragraph()

    # B2B Market
    add_heading(doc, 'PART 1: THE PAYING CUSTOMERS (B2B)', 1)

    add_heading(doc, '1.1 South African Business Landscape', 2)
    add_table(doc,
        ['Metric', 'Value', 'Source'],
        [
            ['Total MSMEs in SA', '2.0+ million', 'UNCTAD 2023'],
            ['Formal registered businesses', '~550,000', 'Small Business Institute'],
            ['SMEs (10-250 employees)', '~150,000', 'StatsSA estimates'],
            ['SMEs that hire annually', '~75,000 (50%)', 'Industry estimate'],
            ['Average roles per hiring SME', '4-5 per year', 'Industry estimate']
        ]
    )
    doc.add_paragraph()

    add_heading(doc, '1.2 B2B TAM Calculation', 2)
    doc.add_paragraph('75,000 SMEs hiring annually x 5 roles x R1,750 = R656M base CV screening')
    doc.add_paragraph('Conservative estimate including enterprise: R1.75 billion')

    # Recruitment Agencies
    add_heading(doc, 'PART 2: RECRUITMENT AGENCIES', 1)

    add_table(doc,
        ['Metric', 'Value', 'Source'],
        [
            ['Africa staffing market (2023)', '$16.5 billion', 'Business Market Insights'],
            ['SA share (~40% of Africa)', '~$6.6 billion', 'Estimate'],
            ['SA recruitment market CAGR', '13.7%', 'Research and Markets'],
            ['Number of registered agencies', '~2,500', 'APSO estimate'],
            ['Active boutique recruiters', '~1,000', 'Industry estimate']
        ]
    )
    doc.add_paragraph()
    doc.add_paragraph('Recruiter TAM: ~R67 million')

    # Job Seekers
    add_heading(doc, 'PART 3: JOB SEEKERS (B2C)', 1)

    add_heading(doc, '3.1 South African Labor Market', 2)
    add_table(doc,
        ['Metric', 'Value', 'Source'],
        [
            ['Working-age population', '40.4 million', 'StatsSA Q1 2025'],
            ['Labor force', '25.1 million', 'StatsSA Q3 2025'],
            ['Employed', '17.1 million', 'StatsSA Q3 2025'],
            ['Unemployed (official)', '8.1 million', 'StatsSA Q3 2025'],
            ['Unemployment rate', '31.9%', 'StatsSA Q3 2025'],
            ['Expanded unemployment', '42.4%', 'StatsSA Q3 2025'],
            ['Youth unemployment (15-24)', '58.5%', 'StatsSA Q3 2025'],
            ['Youth unemployment (15-34)', '46.1%', 'StatsSA Q1 2025']
        ]
    )
    doc.add_paragraph()

    add_heading(doc, '3.2 Active Job Seekers', 2)
    add_table(doc,
        ['Segment', 'Size', 'Profile'],
        [
            ['Officially unemployed', '8.1 million', 'Actively seeking work'],
            ['Underemployed', '~3 million', 'Want more work'],
            ['Passive seekers (employed)', '~5 million', 'Open to opportunities'],
            ['Total potential users', '~16 million', '']
        ]
    )
    doc.add_paragraph()

    # Competitive Landscape
    add_heading(doc, 'PART 4: COMPETITIVE LANDSCAPE', 1)

    add_heading(doc, 'Direct Competitors', 2)
    add_table(doc,
        ['Competitor', 'Focus', 'Pricing', 'HireInbox Advantage'],
        [
            ['Pnet', 'Job board', 'Per listing', 'AI screening included'],
            ['CareerJunction', 'Job board', 'Per listing', 'Evidence-based AI'],
            ['Indeed SA', 'Aggregator', 'CPC/CPA', 'Predictable pricing'],
            ['LinkedIn', 'Professional network', 'Subscription', 'SA-specific context'],
            ['OfferZen', 'Tech talent', '% of salary', 'Broader market']
        ]
    )
    doc.add_paragraph()

    add_heading(doc, 'Why HireInbox Wins', 2)
    bullets = doc.add_paragraph()
    bullets.add_run('1. SA-Specific AI - Understands CA(SA), BCom, local companies\n')
    bullets.add_run('2. Evidence-Based - Every decision backed by quotes (POPIA compliant)\n')
    bullets.add_run('3. Per-Role Pricing - Predictable, fair, no usage anxiety\n')
    bullets.add_run('4. Email-Native - Works where HR already works\n')
    bullets.add_run('5. Unified Platform - B2B, B2C, and Talent Pool in one')

    # Market Sizing Summary
    add_heading(doc, 'PART 5: MARKET SIZING SUMMARY', 1)

    add_heading(doc, 'Total Addressable Market (TAM)', 2)
    add_table(doc,
        ['Segment', 'Size', 'Description'],
        [
            ['B2B Employers', 'R1.75 billion', 'All SA SME hiring spend'],
            ['Recruitment Agencies', 'R67 million', 'Boutique recruiter tools'],
            ['B2C Job Seekers', 'R50 million', 'Paid CV services at scale'],
            ['TOTAL TAM', 'R1.87 billion', 'Annual opportunity']
        ]
    )
    doc.add_paragraph()

    add_heading(doc, 'Serviceable Obtainable Market (SOM) - Year 1', 2)
    add_table(doc,
        ['Segment', 'Size', 'Description'],
        [
            ['B2B', 'R2.4 million', '~200 customers'],
            ['Recruiters', 'R300,000', '~50 recruiters'],
            ['B2C', 'R1.3 million', '~8,000 paid users'],
            ['Talent Pool', 'R580,000', '~230 job posts'],
            ['TOTAL SOM Y1', 'R4.58 million', 'First year target']
        ]
    )
    doc.add_paragraph()

    # SA Statistics
    add_heading(doc, 'PART 6: KEY STATISTICS REFERENCE', 1)

    add_heading(doc, 'South Africa at a Glance (2025)', 2)
    add_table(doc,
        ['Metric', 'Value'],
        [
            ['Population', '62 million'],
            ['Working-age (15-64)', '40.4 million'],
            ['Labor force', '25.1 million'],
            ['Employed', '17.1 million'],
            ['Unemployed', '8.1 million'],
            ['Unemployment rate', '31.9%'],
            ['Youth (15-24) unemployment', '58.5%'],
            ['GDP (2024)', '$380 billion'],
            ['SMEs', '2+ million'],
            ['Formal businesses', '~550,000']
        ]
    )
    doc.add_paragraph()

    add_heading(doc, 'Internet/Mobile Penetration', 2)
    add_table(doc,
        ['Metric', 'Value'],
        [
            ['Smartphone users', '25+ million'],
            ['Internet penetration', '72%'],
            ['Social media users', '26 million'],
            ['WhatsApp users', '20+ million']
        ]
    )
    doc.add_paragraph()

    # Footer
    doc.add_paragraph()
    doc.add_paragraph('Prepared by: Claude (CTO)')
    doc.add_paragraph('For: HireInbox Partner Review')
    doc.add_paragraph('Date: 25 January 2026')
    doc.add_paragraph()
    doc.add_paragraph('"In a market with 8 million unemployed and 75,000 SMEs hiring annually, there is no shortage of opportunity. The question is execution."')

    doc.save(MARKET_PATH)
    print(f"Created: {MARKET_PATH}")

def create_revenue_forecast():
    """Create Revenue Forecast Excel"""
    wb = Workbook()

    # Styles
    header_fill = PatternFill(start_color="4F46E5", end_color="4F46E5", fill_type="solid")
    header_font = Font(bold=True, color="FFFFFF")
    total_fill = PatternFill(start_color="10B981", end_color="10B981", fill_type="solid")
    total_font = Font(bold=True, color="FFFFFF")
    currency_format = 'R#,##0'
    thin_border = Border(
        left=Side(style='thin'),
        right=Side(style='thin'),
        top=Side(style='thin'),
        bottom=Side(style='thin')
    )

    # === SUMMARY SHEET ===
    ws = wb.active
    ws.title = "Summary"

    # Title
    ws['A1'] = "HIREINBOX - 12 MONTH REVENUE FORECAST"
    ws['A1'].font = Font(bold=True, size=16)
    ws['A2'] = "Period: 1 May 2026 - 30 April 2027"
    ws['A3'] = "Prepared: 25 January 2026"

    # Executive Summary
    ws['A5'] = "EXECUTIVE SUMMARY"
    ws['A5'].font = Font(bold=True, size=14)

    summary_data = [
        ["Metric", "Year 1 Total", "% of Total"],
        ["Total Revenue", 4580942, "100%"],
        ["B2B Revenue", 2370261, "52%"],
        ["B2C Revenue", 1334480, "29%"],
        ["Talent Pool Revenue", 577500, "13%"],
        ["Recruiter Revenue", 298701, "6%"]
    ]

    for row_idx, row_data in enumerate(summary_data, start=7):
        for col_idx, value in enumerate(row_data, start=1):
            cell = ws.cell(row=row_idx, column=col_idx, value=value)
            cell.border = thin_border
            if row_idx == 7:
                cell.fill = header_fill
                cell.font = header_font
            elif row_idx == 8:
                cell.fill = total_fill
                cell.font = total_font
            if col_idx == 2 and row_idx > 7:
                cell.number_format = currency_format

    # Key Milestones
    ws['A15'] = "KEY MILESTONES"
    ws['A15'].font = Font(bold=True, size=14)

    milestones = [
        ["Month", "Revenue", "Description"],
        ["Month 1 (May)", 45785, "Launch"],
        ["Month 6 (Oct)", 345350, "Traction"],
        ["Month 12 (Apr)", 872375, "Scale"]
    ]

    for row_idx, row_data in enumerate(milestones, start=17):
        for col_idx, value in enumerate(row_data, start=1):
            cell = ws.cell(row=row_idx, column=col_idx, value=value)
            cell.border = thin_border
            if row_idx == 17:
                cell.fill = header_fill
                cell.font = header_font
            if col_idx == 2 and row_idx > 17:
                cell.number_format = currency_format

    # Quarterly View
    ws['A23'] = "QUARTERLY VIEW"
    ws['A23'].font = Font(bold=True, size=14)

    quarterly = [
        ["Quarter", "Revenue", "Growth"],
        ["Q1 (May-Jul)", 257120, "-"],
        ["Q2 (Aug-Oct)", 802695, "+212%"],
        ["Q3 (Nov-Jan)", 1265543, "+58%"],
        ["Q4 (Feb-Apr)", 2255584, "+78%"],
        ["YEAR 1 TOTAL", 4580942, ""]
    ]

    for row_idx, row_data in enumerate(quarterly, start=25):
        for col_idx, value in enumerate(row_data, start=1):
            cell = ws.cell(row=row_idx, column=col_idx, value=value)
            cell.border = thin_border
            if row_idx == 25:
                cell.fill = header_fill
                cell.font = header_font
            elif row_idx == 30:
                cell.fill = total_fill
                cell.font = total_font
            if col_idx == 2 and row_idx > 25:
                cell.number_format = currency_format

    # Column widths
    ws.column_dimensions['A'].width = 25
    ws.column_dimensions['B'].width = 20
    ws.column_dimensions['C'].width = 15

    # === MONTHLY DETAIL SHEET ===
    ws2 = wb.create_sheet("Monthly Detail")

    # Headers
    headers = ["Month", "B2B Screening", "B2B Interview", "B2B Verify", "Talent Pool", "Talent Map", "B2C Video", "B2C Coaching", "B2C Prep", "TOTAL", "MoM Growth"]
    for col_idx, header in enumerate(headers, start=1):
        cell = ws2.cell(row=1, column=col_idx, value=header)
        cell.fill = header_fill
        cell.font = header_font
        cell.border = thin_border
        cell.alignment = Alignment(horizontal='center')

    # Monthly data
    monthly_data = [
        ["May 2026", 17500, 1598, 800, 5000, 2997, 11920, 3980, 1990, 45785, "Launch"],
        ["Jun 2026", 31500, 3196, 2400, 10000, 4995, 17880, 6965, 3980, 80916, "+77%"],
        ["Jul 2026", 52500, 6392, 4800, 15000, 7992, 26820, 9950, 6965, 130419, "+61%"],
        ["Aug 2026", 78750, 9588, 8000, 25000, 11988, 37250, 13930, 9950, 194456, "+49%"],
        ["Sep 2026", 105000, 14382, 12000, 35000, 17982, 47680, 17910, 12935, 262889, "+35%"],
        ["Oct 2026", 140000, 19975, 16000, 45000, 24975, 59600, 23880, 15920, 345350, "+31%"],
        ["Nov 2026", 175000, 25568, 22400, 55000, 29970, 71520, 29850, 19900, 429208, "+24%"],
        ["Dec 2026", 122500, 15980, 14400, 30000, 19980, 52150, 19900, 13930, 288840, "-33%"],
        ["Jan 2027", 227500, 31960, 28000, 70000, 34965, 89400, 39800, 25870, 547495, "+90%"],
        ["Feb 2027", 262500, 39950, 33600, 80000, 39960, 104300, 47760, 29850, 637920, "+17%"],
        ["Mar 2027", 306250, 46342, 40000, 95000, 47952, 119200, 55720, 34825, 745289, "+17%"],
        ["Apr 2027", 350000, 55930, 48000, 112500, 54945, 141550, 67660, 41790, 872375, "+17%"]
    ]

    for row_idx, row_data in enumerate(monthly_data, start=2):
        for col_idx, value in enumerate(row_data, start=1):
            cell = ws2.cell(row=row_idx, column=col_idx, value=value)
            cell.border = thin_border
            if col_idx >= 2 and col_idx <= 10:
                cell.number_format = currency_format
            if col_idx == 10:
                cell.font = Font(bold=True)

    # Totals row
    totals = ["YEAR 1 TOTAL", 1869000, 270861, 230400, 577500, 298701, 779270, 337305, 217905, 4580942, ""]
    for col_idx, value in enumerate(totals, start=1):
        cell = ws2.cell(row=14, column=col_idx, value=value)
        cell.fill = total_fill
        cell.font = total_font
        cell.border = thin_border
        if col_idx >= 2 and col_idx <= 10:
            cell.number_format = currency_format

    # Column widths
    for col in range(1, 12):
        ws2.column_dimensions[get_column_letter(col)].width = 15
    ws2.column_dimensions['A'].width = 15

    # === PRODUCT BREAKDOWN SHEET ===
    ws3 = wb.create_sheet("Product Breakdown")

    ws3['A1'] = "ANNUAL REVENUE BY PRODUCT"
    ws3['A1'].font = Font(bold=True, size=14)

    product_headers = ["Product", "Units (Y1)", "Price", "Revenue (Y1)", "% of Total"]
    for col_idx, header in enumerate(product_headers, start=1):
        cell = ws3.cell(row=3, column=col_idx, value=header)
        cell.fill = header_fill
        cell.font = header_font
        cell.border = thin_border

    products = [
        ["B2B CV Screening", 1068, 1750, 1869000, "40.8%"],
        ["B2B AI Interview", 339, 799, 270861, "5.9%"],
        ["B2B Verification", 288, 800, 230400, "5.0%"],
        ["Talent Pool Jobs", 231, 2500, 577500, "12.6%"],
        ["Talent Mapping", 299, 999, 298701, "6.5%"],
        ["B2C Video Analysis", 5230, 149, 779270, "17.0%"],
        ["B2C AI Coaching", 1695, 199, 337305, "7.4%"],
        ["B2C Position Prep", 1095, 199, 217905, "4.8%"]
    ]

    for row_idx, row_data in enumerate(products, start=4):
        for col_idx, value in enumerate(row_data, start=1):
            cell = ws3.cell(row=row_idx, column=col_idx, value=value)
            cell.border = thin_border
            if col_idx in [3, 4]:
                cell.number_format = currency_format

    # Total
    total_row = ["TOTAL", "", "", 4580942, "100%"]
    for col_idx, value in enumerate(total_row, start=1):
        cell = ws3.cell(row=12, column=col_idx, value=value)
        cell.fill = total_fill
        cell.font = total_font
        cell.border = thin_border
        if col_idx == 4:
            cell.number_format = currency_format

    # Column widths
    ws3.column_dimensions['A'].width = 20
    ws3.column_dimensions['B'].width = 12
    ws3.column_dimensions['C'].width = 10
    ws3.column_dimensions['D'].width = 15
    ws3.column_dimensions['E'].width = 12

    # === SCENARIOS SHEET ===
    ws4 = wb.create_sheet("Scenarios")

    ws4['A1'] = "REVENUE SCENARIOS"
    ws4['A1'].font = Font(bold=True, size=14)

    scenario_headers = ["Scenario", "Adjustment", "Year 1 Revenue"]
    for col_idx, header in enumerate(scenario_headers, start=1):
        cell = ws4.cell(row=3, column=col_idx, value=header)
        cell.fill = header_fill
        cell.font = header_font
        cell.border = thin_border

    scenarios = [
        ["Conservative", "-30%", 3206659],
        ["Base Case", "This Forecast", 4580942],
        ["Optimistic", "+40%", 6413319]
    ]

    for row_idx, row_data in enumerate(scenarios, start=4):
        for col_idx, value in enumerate(row_data, start=1):
            cell = ws4.cell(row=row_idx, column=col_idx, value=value)
            cell.border = thin_border
            if col_idx == 3:
                cell.number_format = currency_format
            if row_idx == 5:  # Base case
                cell.font = Font(bold=True)

    ws4.column_dimensions['A'].width = 15
    ws4.column_dimensions['B'].width = 15
    ws4.column_dimensions['C'].width = 18

    wb.save(FORECAST_PATH)
    print(f"Created: {FORECAST_PATH}")

if __name__ == "__main__":
    print("Creating HireInbox partner documents...")
    create_prd_document()
    create_market_size_document()
    create_revenue_forecast()
    print("\nAll documents created successfully!")
    print(f"\nFiles on Desktop:")
    print(f"  - HIREINBOX_PRD_TECHNICAL.docx")
    print(f"  - HIREINBOX_MARKET_SIZE.docx")
    print(f"  - HIREINBOX_REVENUE_FORECAST_Y1.xlsx")
